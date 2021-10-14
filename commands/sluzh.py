import command_system
from flask import Flask, request, json
import vkapi
import settings
import keyb
import vk
from PIL import Image, ImageDraw, ImageFont
import sys
from io import BytesIO
import io
import requests
import random
import traceback
import docx



def sluzh():
    data = json.loads(request.data)
    object = data['object']
    message = object['message']
    text = message['text']
    peer_id =  message['peer_id']
    text1 = text.split("\n")
    try:
        phone = text1[8]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        name = text1[7]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        ticket = text1[6]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        end = text1[5]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        start = text1[4]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        date = text1[3]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        cabinet = text1[2]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "c8991d8af3905eaf984126a9ac844c559113e881349d7cdf880ddf8f686044e1153bb60c6af779121add4"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    try:
        event = text1[1]
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    event = text1[1]
    cabinet = text1[2]
    date = text1[3]
    start = text1[4]
    end = text1[5]
    ticket = text1[6]
    name = text1[7]
    phone = text1[8]
    url = 'https://vk.com/doc436020620_584588566?hash=3cfff1000be90ebf3f&dl=GQZTMMBSGA3DEMA:1609959785:a8b11433b429b578f3&api=1&no_preview=1'
    doc = docx.Document(BytesIO(requests.get(url).content))
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("event_name", event)
        paragraph.text = paragraph.text.replace("cabinet_number", cabinet)
        paragraph.text = paragraph.text.replace("event_date", date)
        paragraph.text = paragraph.text.replace("time_start", start)
        paragraph.text = paragraph.text.replace("time_end", end)
        paragraph.text = paragraph.text.replace("ticket_date", ticket)
        paragraph.text = paragraph.text.replace("user_name", name)
        paragraph.text = paragraph.text.replace("user_phone", phone)

    for section in doc.sections:
        footer = section.footer
        footer.paragraphs[0].text  = footer.paragraphs[0].text.replace("user_name", name)
        footer.paragraphs[0].text  = footer.paragraphs[0].text.replace("user_phone", phone)


    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    files={'file': ('report.docx', file_stream)}
    url1 = vkapi.get_doc22(peer_id, settings.access_token)
    response = requests.post(url1 , files=files)
    file_stream.close()
    result = json.loads(response.text)
    file = result['file']
    title = 'Служебка'
    id, owner_id = vkapi.save_doc22(settings.access_token, title, file)
    message = ''
    attachment = 'doc' + str(owner_id) + '_' + str(id)
    keyboard = {"buttons":[],"one_time":True}
    return message, attachment, keyboard


sluzh_command = command_system.Command()

sluzh_command.keys = ['валя служебка']
sluzh_command.process = sluzh
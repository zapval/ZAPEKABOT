import command_system
from flask import Flask, request, json
import vkapi
import settings
import keyb
import vk
from PIL import Image, ImageDraw, ImageFont
import sys
from io import BytesIO
import io
import requests
import random
import traceback
import docx


def otch():
    data = json.loads(request.data)
    object = data['object']
    message = object['message']
    text = message['text']
    peer_id =  message['peer_id']
    name = (text[11:44])
    url = 'https://vk.com/doc436020620_584605258'
    try:
        doc = docx.Document(BytesIO(requests.get(url).content))
    except:
        peer_id =  message['peer_id']
        message = traceback.format_exc()
        attachment = ''
        random_id = random.randint(1,100)
        token = "e4b9fa2affa7571dd84575f65a9159830f14533b6c2cdf7f981794795b22fb54c1b7b691b845468fd9c27"
        attachment = ''
        keyboard = keyboard = {"buttons":[],"one_time":True}
        keyboard = json.dumps(keyboard, ensure_ascii = False)
        vkapi.send_message(peer_id, token, keyboard, random_id, message, attachment)
    doc = docx.Document(BytesIO(requests.get(url).content))
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("name", name)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    files={'file': ('report.docx', file_stream)}
    url1 = vkapi.get_doc22(peer_id, settings.access_token)
    response = requests.post(url1 , files=files)
    file_stream.close()
    result = json.loads(response.text)
    file = result['file']
    title = 'Приказ на отчисление ' + name
    id, owner_id = vkapi.save_doc22(settings.access_token, title, file)
    message = 'ну привет, ' + name
    attachment = 'doc' + str(owner_id) + '_' + str(id)
    keyboard = {"buttons":[],"one_time":True}
    return message, attachment, keyboard
    
    
otch_command = command_system.Command()

otch_command.keys = ['отчисление']
otch_command.process = otch
